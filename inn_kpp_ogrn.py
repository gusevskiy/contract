from pathlib import Path
import re

from docx import Document


class Openfile:
    def __init__(self, file_path):
        if file_path.endswith('.pdf'):
            file_path = Path(file_path).stem + ".docx"
            print(file_path)
            self.file_path = file_path
            self.file_data = None
        else:
            self.file_path = file_path
            self.file_data = None

    def reed_inn(self):
        inn = r"[А-Яа-я]{3}[-|:| ]+(\d{10})"
        self.file_data = Document(self.file_path)
        all_par = self.file_data.paragraphs
        for i in all_par:
            if re.findall(inn, i.text):
                return re.findall(inn, i.text)

    def reed_kpp(self):
        kpp = r"[А-Яа-я]{3}[:| ]+(\d{9})"
        self.file_data = Document(self.file_path)
        all_par = self.file_data.paragraphs
        for i in all_par:
            if re.findall(kpp, i.text):
                return re.findall(kpp, i.text)

    def reed_ogrn(self):
        ogrn = r"[ОГРНогрн{4}[-|:| ]+\d{13}"
        self.file_data = Document(self.file_path)
        all_par = self.file_data.paragraphs
        for i in all_par:
            if re.findall(ogrn, i.text):
                return re.findall(ogrn, i.text)


if __name__ == '__main__':
    data_recvi = Openfile('test_class/spravka.pdf')
    # print(data_recvi)
    print(data_recvi.reed_inn())
    print(data_recvi.reed_kpp())
    print(data_recvi.reed_ogrn())
    # data_spravka = Openfile('spravka.docx')
    # data_spravka.reed()
